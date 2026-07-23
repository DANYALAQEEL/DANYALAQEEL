from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def create_professional_docx():
    doc = Document()

    # --- Styles Setup ---
    # Modify basic styles for a cleaner look
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(46, 116, 181) # Corporate Blue
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri Light'
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(46, 116, 181)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    
    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(68, 84, 106) # Dark Gray
    h3.font.bold = True

    # Title Style
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri Light'
    title_style.font.size = Pt(26)
    title_style.font.color.rgb = RGBColor(31, 78, 121) 

    # --- Cover Page ---
    doc.add_paragraph().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Spacing
    
    title = doc.add_paragraph('Software Design Specification (SDS)', style='Title')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('An Intelligent Framework for Automated Data Extraction from Scientific Charts')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.style = 'Normal'
    subtitle.runs[0].font.name = 'Calibri Light'
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    
    # Spacing
    for _ in range(5):
        doc.add_paragraph()

    # Team Info
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.autofit = True
    
    # We'll just put team info in a clean paragraph block instead of a complex table for simplicity on the cover
    team_p = doc.add_paragraph()
    team_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    team_run = team_p.add_run('Team Members:\n')
    team_run.bold = True
    team_run.font.size = Pt(12)
    
    team_details = [
        "Danyal Aqeel (503823)",
        "Muhammad Abdullah Rana (508388)",
        "Huzaifa Sohail (508652)",
        "Annan Khan (506644)"
    ]
    for member in team_details:
        team_p.add_run(f'{member}\n')

    doc.add_paragraph().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Spacing

    org_p = doc.add_paragraph()
    org_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    org_p.add_run('Organization:\n').bold = True
    org_p.add_run('Dept of Computing, SEECS, NUST\n\n')
    
    org_p.add_run('Client:\n').bold = True
    org_p.add_run('Victreat (NSTP)\n\n')
    
    org_p.add_run('Date:\n').bold = True
    org_p.add_run('December 12, 2025')

    doc.add_page_break()

    # --- Content ---

    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph('The purpose of this Software Design Specification (SDS) is to detail the architectural design, system modules, data flow, and design decisions for the "Chart-to-Table Extraction" system. It serves as a blueprint for the development team, ensuring a clear understanding of the system’s structure and behavior before implementation. This document bridges the gap between the requirements analysis and the actual coding phase.')

    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph('The system is a web-based application designed to automate the extraction of underlying data from static scientific chart images (e.g., bar charts, line graphs). Utilizing a hybrid pipeline that combines Deep Learning (DePlot Transformer model) and Computer Vision (Tesseract OCR, OpenCV), the system converts visual chart data into structured formats such as CSV or JSON. The solution includes a React-based frontend for user interaction and a Flask-based backend for processing.')

    doc.add_heading('1.3 Definitions, Acronyms, and Abbreviations', level=2)
    p = doc.add_paragraph()
    p.add_run('OCR (Optical Character Recognition):').bold = True
    p.add_run(' Technology used to distinguish printed or handwritten text characters inside digital images of physical documents.\n')
    p.add_run('CLAHE (Contrast Limited Adaptive Histogram Equalization):').bold = True
    p.add_run(' A computer vision technique used to improve the contrast in images.\n')
    p.add_run('API (Application Programming Interface):').bold = True
    p.add_run(' A set of functions and procedures allowing the creation of applications that access the features or data of an operating system, application, or other service.\n')
    p.add_run('DePlot:').bold = True
    p.add_run(' A state-of-the-art visual language model designed specifically for chart-to-table translation.\n')
    p.add_run('SDS:').bold = True
    p.add_run(' Software Design Specification.\n')
    p.add_run('JSON:').bold = True
    p.add_run(' JavaScript Object Notation.\n')
    p.add_run('CSV:').bold = True
    p.add_run(' Comma-Separated Values.')

    doc.add_heading('1.4 References', level=2)
    p = doc.add_paragraph(style='List Number')
    p.add_run('Assignment 2 - Software Requirements Specification (SRS).')
    p = doc.add_paragraph(style='List Number')
    p.add_run('IEEE 1016-2009 Standard for Information Technology – Systems Design – Software Design Descriptions.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Victreat Project Proposal.')

    doc.add_heading('2. System Overview', level=1)
    
    doc.add_heading('2.1 Description', level=2)
    doc.add_paragraph('The "Chart-to-Table Extraction" system is a client-server application. The workflow is initiated by the user uploading an image of a chart via the web frontend. The backend server receives this image and processes it through a multi-stage pipeline involving preprocessing, text generation (DePlot), and OCR-based validation. The extracted data is then returned to the frontend, where it is displayed to the user along with a confidence score. Users can edit the data if necessary and export it.')

    doc.add_heading('2.2 Major Features', level=2)
    features = [
        ("Image Preprocessing", "Enhances input images using Denoising and Resizing techniques to improve model accuracy."),
        ("Chart Linearization", "Utilizes the DePlot model to convert chart features into a linearized text representation."),
        ("OCR Validation", "Employs Tesseract OCR to cross-verify the data generated by the model against text detected directly from the image axes."),
        ("Data Export", "Allows users to download the extracted data in CSV or JSON formats."),
        ("History Management", "Stores previous extraction results locally for quick access.")
    ]
    for feat, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{feat}: ').bold = True
        p.add_run(desc)

    doc.add_heading('2.3 Technology Stack', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Frontend: ').bold = True
    p.add_run('React, Vite, Tailwind CSS.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Backend: ').bold = True
    p.add_run('Python, Flask, PyTorch, Transformers, OpenCV, Tesseract.')

    doc.add_heading('3. Architectural Design', level=1)
    
    doc.add_heading('3.1 Architecture Pattern', level=2)
    doc.add_paragraph('The system follows a Layered Architecture. This promotes separation of concerns and maintainability.')
    layers = [
        ("Presentation Layer", "The client-side application (React) handling user interactions."),
        ("Application Logic Layer", "The Flask server and Pipeline Controller that manage requests and orchestrate the extraction process."),
        ("Data Processing Layer", "The core AI/ML modules including Preprocessor, DePlot Model, OCREngine, and Validator.")
    ]
    for layer, desc in layers:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{layer}: ').bold = True
        p.add_run(desc)

    doc.add_heading('3.2 Subsystems', level=2)
    subsystems = [
        ("Frontend Subsystem", "Responsible for the User Interface, handling file uploads, displaying results, and managing local state."),
        ("API Gateway", "The Flask web server that exposes REST endpoints for the frontend to communicate with."),
        ("Extraction Pipeline", "The core logic engine that encapsulates the AI models and image processing algorithms.")
    ]
    for sub, desc in subsystems:
        p = doc.add_paragraph()
        p.add_run(f'{sub}: ').bold = True
        p.add_run(desc)

    doc.add_heading('3.3 Diagram Description', level=2)
    doc.add_paragraph('The high-level data flow is as follows:')
    steps = [
        "Client: User uploads an image.",
        "Flask Server: Receives the POST request.",
        "Pipeline Controller: Instantiates the pipeline.",
        "Preprocessor: Applies denoising and CLAHE to the image.",
        "DePlot Model: Generates raw linearized text/table from the image.",
        "OCR Engine: Extracts text ticks from axes for validation.",
        "Validator: Compares Model output with OCR output to compute a confidence score.",
        "Postprocessor: Structures the data into DataFrame/JSON.",
        "Response: JSON data returned to Client."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('4. Detailed Design', level=1)
    
    doc.add_heading('4.1 Component/Class Design', level=2)
    
    doc.add_heading('4.1.1 Class ChartExtractionPipeline', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Attributes: ').bold = True
    p.add_run('preprocessor, model, ocr, validator.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Method: ').bold = True
    p.add_run('process_image(image_path) - Orchestrates the flow by calling sub-components in order.')

    doc.add_heading('4.1.2 Class Preprocessor', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Method: ').bold = True
    p.add_run('preprocess(image)')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Logic: ').bold = True
    p.add_run('Uses cv2.fastNlMeansDenoisingColored to remove noise and cv2.createCLAHE to enhance local contrast, preparing the image for the model.')

    doc.add_heading('4.1.3 Class DePlotModel', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Description: ').bold = True
    p.add_run('Wrapper around the HuggingFace Pix2StructProcessor and model.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Method: ').bold = True
    p.add_run('generate_linearized_text(image) - Passes the processed image to the Transformer model to obtain the textual representation of the table.')

    doc.add_heading('4.1.4 Class OCREngine', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Method: ').bold = True
    p.add_run('extract_text(image)')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Logic: ').bold = True
    p.add_run('Uses pytesseract to detect numbers and labels on the chart axes.')

    doc.add_heading('4.1.5 Class Validator', level=3)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Method: ').bold = True
    p.add_run('validate(df, ocr_ticks)')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Logic: ').bold = True
    p.add_run('Compares the numerical range and values from the DePlot DataFrame (df) against the ticks detected by OCR (ocr_ticks) to determine if the extraction is hallucinated or accurate. Returns a confidence score.')

    doc.add_heading('4.1.6 Frontend Components', level=3)
    components = [
        ("Upload.jsx", "Drag-and-drop zone for file selection and preview."),
        ("Results.jsx", "Displays the extracted table (editable) and the calculated confidence score."),
        ("History.jsx", "Lists past extractions stored in local storage.")
    ]
    for comp, desc in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{comp}: ').bold = True
        p.add_run(desc)

    doc.add_heading('4.2 Interface Design', level=2)
    doc.add_heading('4.2.1 API Endpoints', level=3)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Endpoint: ').bold = True
    p.add_run('POST /api/convert')
    
    sub_p = doc.add_paragraph(style='List Bullet 2')
    sub_p.add_run('Input: ').bold = True
    sub_p.add_run('multipart/form-data containing the image file.')
    
    sub_p = doc.add_paragraph(style='List Bullet 2')
    sub_p.add_run('Output: ').bold = True
    sub_p.add_run('JSON object:')
    
    code = """{
  "success": true, 
  "data": [ ... ],
  "columns": [ ... ],
  "accuracy": { "score": 0.95 }
}"""
    p = doc.add_paragraph(code)
    p.style = 'Normal'
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Endpoint: ').bold = True
    p.add_run('POST /api/convert/csv')
    
    sub_p = doc.add_paragraph(style='List Bullet 2')
    sub_p.add_run('Returns: ').bold = True
    sub_p.add_run('A downloadable .csv file of the extracted data.')

    doc.add_heading('4.3 Data Design', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Backend Storage: ').bold = True
    p.add_run('No persistent database. Temporary storage (temp_outputs/) is used for table.json and table.csv during the session.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Frontend Storage: ').bold = True
    p.add_run('Uses browser localStorage under the key "chart_extraction_history" to store metadata (filename, timestamp, confidence) and a snippet of the data.')

    doc.add_heading('4.4 Behavior Modeling', level=2)
    doc.add_paragraph('Sequence of Operations:')
    seq = [
        "User Uploads file via Frontend.",
        "Frontend sends asynchronous POST request to Backend.",
        "Backend saves file momentarily.",
        "Pipeline runs:\n    a. Preprocessor enhances image.\n    b. Model generates linearized text.\n    c. Postprocessor converts text to DataFrame.\n    d. Validator cross-references DataFrame values with OCR results.",
        "Result JSON is sent back to Frontend.",
        "Frontend renders the Table data."
    ]
    for step in seq:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('4.5 Algorithm Design (Pseudocode)', level=2)
    doc.add_paragraph('Validation Algorithm (from validator.py):')
    
    pseudocode = """FUNCTION Validate(dataframe, ocr_results):
    IF dataframe is EMPTY:
        RETURN confidence = 0

    IF dataframe contains NULL values:
        REDUCE confidence

    max_value = EXTRACT max value from dataframe
    max_tick = EXTRACT max number from ocr_results

    # Calculate agreement ratio
    ratio = max_value / max_tick

    confidence = 0.5 (Base)

    IF ratio is approx 1.0:
        INCREASE confidence (Successful Match)
    ELSE IF ratio indicates unit mismatch (e.g., 1000 vs 1k):
        FLAG warning
        ADJUST confidence based on heuristics

    RETURN confidence"""
    p = doc.add_paragraph(pseudocode)
    p.style = 'Normal'
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('5. User Interface Design', level=1)
    
    doc.add_heading('5.1 Screens', level=2)
    screens = [
        ("Home", "Features a Hero section with a clear 'Start Extraction' call-to-action."),
        ("Upload", "Contains a drag-and-drop zone with client-side file type validation (accepts PNG, JPG)."),
        ("Results", "A split-view layout. The left side shows the extracted data in a grid/table. The right side displays the Verification panel with the confidence score and warnings. Cells in the table are editable."),
        ("History", "A grid view displaying cards of past extractions.")
    ]
    for scr, desc in screens:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{scr}: ').bold = True
        p.add_run(desc)

    doc.add_heading('5.2 Navigation', level=2)
    doc.add_paragraph('The application uses React Router for client-side routing:')
    navs = [
        "/ (Home)",
        "/upload (Upload Interface)",
        "/results (display extraction output)",
        "/history (View past logs)"
    ]
    for nav in navs:
        doc.add_paragraph(nav, style='List Bullet')

    doc.add_heading('6. External Interface Design', level=1)
    
    doc.add_heading('6.1 Software Interfaces', level=2)
    softs = [
        ("Python Environment", "Requires Python 3.8+ with specific library versions defined in requirements.txt."),
        ("Node.js Environment", "Required for building and serving the React development server."),
        ("Tesseract OCR", "The binary tesseract must be installed and accessible via system PATH or configured path.")
    ]
    for soft, desc in softs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{soft}: ').bold = True
        p.add_run(desc)

    doc.add_heading('6.2 Communication Protocols', level=2)
    doc.add_paragraph('Communication happens via REST over HTTP.')
    prots = [
        ("Flask Backend", "Runs on Port 5000 (default)."),
        ("Vite Frontend", "Runs on Port 5173 (default).")
    ]
    for prot, desc in prots:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{prot}: ').bold = True
        p.add_run(desc)

    doc.add_heading('7. Non-Functional Requirements', level=1)
    
    reqs = [
        ("7.1 Performance", [
            ("Inference Latency", "The system relies on a Transformer model; however, the target inference time is under 15 seconds per image."),
            ("Hardware Acceleration", "The system should utilize GPU (CUDA) if available to speed up DePlot and PyTorch operations.")
        ]),
        ("7.2 Reliability", [
            ("Input Validation", "Backend must reject unsupported file types or corrupted images."),
            ("Failure Handling", "If OCR fails to detect text, the system should fall back to just the Model output, marking the confidence as 'Unverified' rather than crashing.")
        ]),
        ("7.3 Maintainability", [
            ("Modularity", "The codebase follows a Pipeline pattern, allowing individual modules (e.g., swapping OCREngine) to be upgraded without affecting the rest of the system.")
        ])
    ]

    for title, items in reqs:
        doc.add_heading(title, level=2)
        for sub_title, text in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'{sub_title}: ').bold = True
            p.add_run(text)

    doc.add_heading('8. Design Constraints', level=1)
    constraints = [
        ("Model Availability", "The system depends on the availability of the pre-trained DePlot model from HuggingFace Hub."),
        ("System Dependencies", "Requires specifically Tesseract OCR to be installed on the host machine; it is not a pure Python dependency.")
    ]
    for cons, desc in constraints:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{cons}: ').bold = True
        p.add_run(desc)

    doc.add_heading('9. Assumptions & Dependencies', level=1)
    assumptions = [
        ("Assumption", "Input images are of reasonable quality and resolution where text is legible to the human eye."),
        ("Assumption", "Charts are standard types (Bar, Line, Scatter) with visible axes."),
        ("Dependency", "transformers, torch, opencv-python, pytesseract libraries.")
    ]
    for type_, desc in assumptions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{type_}: ').bold = True
        p.add_run(desc)

    doc.add_heading('10. Appendices', level=1)
    doc.add_heading('10.1 Glossary', level=2)
    glossary = [
        ("JSON", "JavaScript Object Notation - an open standard file format and data interchange format."),
        ("CSV", "Comma-Separated Values - a delimited text file that uses a comma to separate values."),
        ("SDS", "Software Design Specification."),
        ("REST", "Representational State Transfer - a software architectural style for creating Web services.")
    ]
    for term, definition in glossary:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{term}: ').bold = True
        p.add_run(definition)

    # Save
    file_path = "C:\\Users\\Administrator\\.gemini\\antigravity\\scratch\\Professional_SDS.docx"
    doc.save(file_path)
    print(f"Document saved to {file_path}")

if __name__ == "__main__":
    create_professional_docx()
