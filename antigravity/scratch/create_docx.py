import docx
from docx.shared import Pt, Inches

doc = docx.Document()

with open(r'C:\Users\Administrator\.gemini\antigravity\scratch\rewritten_strategy_report.md', 'r', encoding='utf-8') as f:
    for line in f:
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

doc.save(r'C:\Users\Administrator\Downloads\SDA-Assignment1-Strategy-Rewritten.docx')
print("Document saved successfully.")
